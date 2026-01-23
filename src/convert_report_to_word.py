
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(input_md, output_docx):
    print(f"Converting {input_md} to {output_docx}...")
    
    doc = Document()
    
    # Title Style
    style_title = doc.styles['Title']
    style_title.font.name = 'Malgun Gothic'
    style_title.font.size = Pt(24)
    style_title.font.bold = True
    
    # Heading 1 Style
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Malgun Gothic'
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 Style
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Malgun Gothic'
    style_h2.font.size = Pt(14)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    
    # Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Malgun Gothic'
    style_normal.font.size = Pt(11)
    
    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Simple Markdown Parser
    in_table = False
    table_header = []
    table_rows = []
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines if not in table
        if not line:
            if in_table:
                # Flush table
                _add_table(doc, table_header, table_rows)
                in_table = False
                table_header = []
                table_rows = []
            continue
            
        # Headers
        if line.startswith('# '):
            if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
            doc.add_heading(line[4:], level=2)
            
        # List items
        elif line.startswith('- ') or line.startswith('* '):
            if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
            
        # Numbered list
        elif line[0].isdigit() and line[1] == '.':
             if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
             p = doc.add_paragraph(style='List Number')
             _add_formatted_text(p, line.split('.', 1)[1].strip())
             
        # Table
        elif line.startswith('|'):
            if '---' in line:
                continue # Skip separator
                
            cells = [c.strip() for c in line.strip('|').split('|')]
            
            if not in_table:
                in_table = True
                table_header = cells
                table_rows = []
            else:
                table_rows.append(cells)
                
        # Normal text
        else:
            if in_table:
                # Check if this line is actually part of table or break
                # Assuming table lines must start with |
                _add_table(doc, table_header, table_rows)
                in_table = False
                
            if line.startswith('---'):
                doc.add_page_break()
                continue
                
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            
    # Flush remaining table
    if in_table:
        _add_table(doc, table_header, table_rows)
        
    doc.save(output_docx)
    print(f"Saved: {output_docx}")

def _add_formatted_text(paragraph, text):
    # Handle simple bolding **text**
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1: # Odd parts are bold
            run.bold = True
            
def _add_table(doc, header, rows):
    if not header: return
    
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    # Rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i < len(row_cells):
                # Handle formatted text in cells too
                # For simplicity, just strips
                if '**' in cell_text:
                    p = row_cells[i].paragraphs[0]
                    _add_formatted_text(p, cell_text)
                else:
                    row_cells[i].text = cell_text

if __name__ == '__main__':
    create_word_report('/data/ethylene/FINAL_PROJECT_REPORT.md', '/data/ethylene/FINAL_PROJECT_REPORT.docx')
