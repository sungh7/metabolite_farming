
import os
import re
import csv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report(input_md, output_docx):
    print(f"Converting {input_md} to {output_docx} (with linked files)...")
    
    doc = Document()
    
    # Styles Setup (Malgun Gothic for consistency)
    _setup_styles(doc)
    
    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    # Track linked files to append later
    linked_files = [] 
    
    # Simple Markdown Parser
    in_table = False
    table_header = []
    table_rows = []
    
    for line in lines:
        line = line.strip()
        
        # Detect Links: [Label](Path)
        # We find all links, check if they are local files, and store for appendix
        links = re.findall(r'\[(.*?)\]\((.*?)\)', line)
        for label, path in links:
            if not path.startswith('http') and not path.startswith('#'):
                # Handle relative paths properly
                abs_path = path
                if not os.path.isabs(path):
                    # Relative to input markdown file location
                    base_dir = os.path.dirname(os.path.abspath(input_md))
                    abs_path = os.path.join(base_dir, path)
                
                # Check exist
                if os.path.exists(abs_path):
                    # Avoid duplicates
                    if abs_path not in [x['path'] for x in linked_files]:
                         linked_files.append({'label': label, 'path': abs_path})

        
        # Skip empty lines if not in table
        if not line:
            if in_table:
                _add_table(doc, table_header, table_rows)
                in_table = False
                table_header = []
                table_rows = []
            continue
            
        # Headers
        if line.startswith('# '):
            if in_table: _add_table(doc, table_header, table_rows); in_table = False
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            if in_table: _add_table(doc, table_header, table_rows); in_table = False
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            if in_table: _add_table(doc, table_header, table_rows); in_table = False
            doc.add_heading(line[4:], level=2)
            
        # List items
        elif line.startswith('- ') or line.startswith('* '):
            if in_table: _add_table(doc, table_header, table_rows); in_table = False
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
            
        # Numbered list
        elif line[0].isdigit() and line[1] == '.':
             if in_table: _add_table(doc, table_header, table_rows); in_table = False
             p = doc.add_paragraph(style='List Number')
             _add_formatted_text(p, line.split('.', 1)[1].strip())
             
        # Table
        elif line.startswith('|'):
            if '---' in line: continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            if not in_table:
                in_table = True; table_header = cells; table_rows = []
            else:
                table_rows.append(cells)
                
        # Normal text
        else:
            if in_table: _add_table(doc, table_header, table_rows); in_table = False
            if line.startswith('---'): doc.add_page_break(); continue
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
            
    if in_table:
        _add_table(doc, table_header, table_rows)
        
    # ==========================================
    # Append Linked Files Content
    # ==========================================
    if linked_files:
        doc.add_page_break()
        doc.add_heading('Appendix: Linked Files Content', level=0)
        
        for item in linked_files:
            label = item['label']
            path = item['path']
            filename = os.path.basename(path)
            ext = os.path.splitext(filename)[1].lower()
            
            doc.add_heading(f"{label} ({filename})", level=1)
            p = doc.add_paragraph()
            p.add_run(f"FilePath: {path}").italic = True
            
            try:
                if ext in ['.png', '.jpg', '.jpeg']:
                    doc.add_picture(path, width=Inches(6.0))
                    
                elif ext == '.csv':
                    with open(path, 'r', encoding='utf-8') as csvfile:
                        reader = csv.reader(csvfile)
                        headers = next(reader, None)
                        rows = list(reader)
                        if headers:
                            _add_table(doc, headers, rows)
                        else:
                            doc.add_paragraph("(Empty CSV)")
                            
                else:
                    # Text files (md, txt, py, etc.)
                    with open(path, 'r', encoding='utf-8', errors='replace') as tf:
                        content = tf.read()
                        
                    # Add as code block-ish text (Monospace)
                    p = doc.add_paragraph()
                    run = p.add_run(content[:5000]) # Retrieve first 5000 chars to avoid overload
                    if len(content) > 5000:
                        run.text += "\n... (Truncated)"
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    
            except Exception as e:
                doc.add_paragraph(f"Error reading file: {str(e)}")
                
            doc.add_paragraph() # Spacing
        
    doc.save(output_docx)
    print(f"Saved: {output_docx}")

def _setup_styles(doc):
    style_title = doc.styles['Title']
    style_title.font.name = 'Malgun Gothic'
    style_title.font.size = Pt(24)
    style_title.font.bold = True
    
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Malgun Gothic'
    style_h1.font.size = Pt(16)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Malgun Gothic'
    style_normal.font.size = Pt(11)

def _add_formatted_text(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1: run.bold = True
            
def _add_table(doc, header, rows):
    if not header: return
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            if i < len(row_cells):
                 row_cells[i].text = cell_text

if __name__ == '__main__':
    create_word_report('/data/ethylene/FINAL_PROJECT_REPORT.md', '/data/ethylene/FINAL_PROJECT_REPORT.docx')
